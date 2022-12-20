import os
import sys
import docx
import log

NAME = 'result'


def average_temperature(temp_l: list[str]) -> float:
    """
    Вычисляем среднюю температуру
    :param temp_l: Значения температуры, строковые -> [str(float), '37.1', '37.2', ... ],
    len(temp_list) = 9 (matrix[3*3])
    :return: Средняя температура : float
    """
    t_sr = 0
    for s in temp_l:
        t_sr = t_sr + float(s)
    return t_sr / len(temp_l)


def add_result(uid: str, res_str: str, temp_list: list, name:str = NAME) -> None:
    """
    Заполняем созданный ранее {NAME}.docx
    temp_list -> [str(float), '37.1', '37.2', ... ], len(temp_list) = 9 (matrix[3*3])
    """
    document_path = f"{os.path.dirname(os.path.realpath(__file__))}{os.sep}{name}.docx"
    document = docx.Document(document_path)
    document.add_heading(uid, level=1)

    # Проверка, что не сдвинуты занчения, если сдвинуты, не обрабатываем их
    if temp_list[0] == '':
        print(f"{uid} is shifted: {temp_list}")
        document.add_paragraph('The value in the database is shifted, this study is not processed')
        document.save(document_path)
        log.add_to_log(f"{uid} is shifted: {temp_list}")
        return None

    if temp_list[-1] == 0:
        print(f"{uid} is shifted: {temp_list}")
        document.add_paragraph('The value in the database is shifted, this study is not processed')
        document.save(document_path)
        log.add_to_log(f"{uid} is shifted: {temp_list}")
        return None

    # Добавляем таблицу
    table = document.add_table(rows=1, cols=10)
    hdr_cells = table.rows[0].cells

    # создали таблицу и заполнили ее значениями температуры (внутренеей, не температурой кожи)
    hdr_cells[0].text = 'Internal Temp'

    for i, t in enumerate(temp_list):
        hdr_cells[i + 1].text = t

    # Добавили среднее значение температуры
    row_cells = table.add_row().cells
    row_cells[0].text = 'Average Temp'
    a_temp = average_temperature(temp_list)
    for i in range(9):
        row_cells[i+1].text = str(round(a_temp, 2))

    # Заполнили таблицу разницей между внутренней и средним значением
    r_list = []
    row_cells = table.add_row().cells
    row_cells[0].text = 'Delta Temp'
    for i, temp in enumerate(temp_list):
        row_cells[i + 1].text = str(round(float(temp) - a_temp,2))
        # Если дельта >=1 или >=0.5 создаем "флаг"
        if float(temp) - a_temp >= 1:
            r_list.append('C')
        elif float(temp) - a_temp >= 0.5:
            r_list.append('M')
        else:
            r_list.append('n')

    # Записываем результат согласно "флагу"
    if 'C' in r_list:
        document.add_paragraph('Critical delta >= 1.0')
    elif 'M' in r_list:
        document.add_paragraph("Warning 0.5 <= delta < 1.0")
    else:
        document.add_paragraph('Norma delta < 0.5')
    # добавляем из базы диагноз врача
    document.add_paragraph(res_str)
    # сохраняем файл
    document.save(document_path)
    return None


def create_empty_docx(base_path: str, name:str = NAME) -> None:
    """
    Создаем пустой файл
    :param base_path: Путь к базе данных, будет записан в заголовке
    :param name: Имя файла для сохранения, по умолчанию result
    :return: None
    """
    try:
        document = docx.Document()
        document.add_heading(base_path, 0)
        document.save(f"{os.path.dirname(os.path.realpath(__file__))}{os.sep}{name}.docx")
    except Exception as err:
        log.add_warn(f"Error, can`t create docx file: {err}")
    return None



